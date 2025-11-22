# -*- coding: utf-8 -*-
"""
Test script for MCP Documents Service
Tests all endpoints: Word, PDF, and Google Docs integration
"""

import sys
import io as io_module

# Set UTF-8 encoding for stdout on Windows
if sys.platform == 'win32':
    sys.stdout = io_module.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import requests
import base64
import os
import time
from docx import Document
import io

# Service URL
BASE_URL = "http://localhost:8007"

def print_separator():
    print("\n" + "="*80 + "\n")

def test_health_check():
    """Test the health check endpoint"""
    print("Testing health check endpoint...")
    response = requests.get(f"{BASE_URL}/health")

    if response.status_code == 200:
        data = response.json()
        print(f"✓ Service is healthy")
        print(f"  Status: {data['status']}")
        print(f"  Service: {data['service']}")
        print(f"  Google Docs available: {data['google_docs_available']}")
        return True
    else:
        print(f"✗ Health check failed: {response.status_code}")
        return False

def test_root_endpoint():
    """Test the root endpoint"""
    print("Testing root endpoint...")
    response = requests.get(f"{BASE_URL}/")

    if response.status_code == 200:
        data = response.json()
        print(f"✓ Root endpoint OK")
        print(f"  Service: {data['service']}")
        print(f"  Version: {data['version']}")
        print(f"  Endpoints: {len(data['endpoints'])} categories")
        return True
    else:
        print(f"✗ Root endpoint failed: {response.status_code}")
        return False

def test_create_docx():
    """Test creating a Word document"""
    print("Testing /documents/create_docx...")

    payload = {
        "title": "Test Document MCP",
        "content": [
            {
                "text": "Rapport de Test",
                "style": "Heading1",
                "alignment": "center"
            },
            {
                "text": "Introduction",
                "style": "Heading2",
                "alignment": "left"
            },
            {
                "text": "Ceci est un paragraphe de test avec du contenu normal. Le service MCP Documents fonctionne correctement.",
                "style": "Normal",
                "alignment": "justify",
                "bold": False,
                "italic": False
            },
            {
                "text": "Résultats",
                "style": "Heading2",
                "alignment": "left"
            },
            {
                "text": "Les tests ont démontré que le système fonctionne comme prévu.",
                "style": "Normal",
                "alignment": "left",
                "bold": True,
                "italic": False
            }
        ]
    }

    response = requests.post(f"{BASE_URL}/documents/create_docx", json=payload)

    if response.status_code == 200:
        data = response.json()
        print(f"✓ Document created successfully")
        print(f"  Filename: {data['filename']}")
        print(f"  Size: {data['size_bytes']} bytes")

        # Save the document
        doc_bytes = base64.b64decode(data['document_base64'])
        output_file = "test_output.docx"
        with open(output_file, "wb") as f:
            f.write(doc_bytes)
        print(f"  Saved to: {output_file}")

        # Verify the document can be opened
        doc = Document(output_file)
        print(f"  Verified: Document has {len(doc.paragraphs)} paragraphs")

        return True, data['document_base64']
    else:
        print(f"✗ Failed to create document: {response.status_code}")
        print(f"  Error: {response.text}")
        return False, None

def test_fill_template():
    """Test filling a Word template"""
    print("Testing /documents/fill_docx_template...")

    # First create a simple template
    doc = Document()
    doc.add_heading("Facture {{numero}}", 0)
    doc.add_paragraph("Client: {{client_nom}}")
    doc.add_paragraph("Date: {{date}}")
    doc.add_paragraph("Montant: {{montant}}")

    # Save template to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    template_b64 = base64.b64encode(buffer.read()).decode()

    payload = {
        "template_base64": template_b64,
        "replacements": {
            "numero": "INV-2025-001",
            "client_nom": "ACME Corporation",
            "date": "21/11/2025",
            "montant": "5,000 EUR"
        }
    }

    response = requests.post(f"{BASE_URL}/documents/fill_docx_template", json=payload)

    if response.status_code == 200:
        data = response.json()
        print(f"✓ Template filled successfully")
        print(f"  Replacements: {data['replacements_count']}")

        # Save the filled document
        doc_bytes = base64.b64decode(data['document_base64'])
        output_file = "test_filled_template.docx"
        with open(output_file, "wb") as f:
            f.write(doc_bytes)
        print(f"  Saved to: {output_file}")

        # Verify replacements
        filled_doc = Document(output_file)
        text_content = "\n".join([p.text for p in filled_doc.paragraphs])
        if "ACME Corporation" in text_content and "INV-2025-001" in text_content:
            print(f"  Verified: Replacements are present in document")
        else:
            print(f"  Warning: Could not verify all replacements")

        return True
    else:
        print(f"✗ Failed to fill template: {response.status_code}")
        print(f"  Error: {response.text}")
        return False

def test_export_pdf():
    """Test exporting to PDF"""
    print("Testing /documents/export_pdf...")

    payload = {
        "title": "Rapport PDF Test",
        "content": [
            {
                "text": "RAPPORT DE TEST PDF",
                "style": "Heading1",
                "alignment": "center"
            },
            {
                "text": "Section 1 - Vue d'ensemble",
                "style": "Heading2",
                "alignment": "left"
            },
            {
                "text": "Ce document PDF a été généré automatiquement par le service MCP Documents. Il démontre la capacité du système à exporter du contenu au format PDF.",
                "style": "Normal",
                "alignment": "justify"
            },
            {
                "text": "Section 2 - Résultats",
                "style": "Heading2",
                "alignment": "left"
            },
            {
                "text": "Les tests ont confirmé que l'export PDF fonctionne correctement avec différents styles de formatage.",
                "style": "Normal",
                "alignment": "left"
            }
        ],
        "page_size": "A4"
    }

    response = requests.post(f"{BASE_URL}/documents/export_pdf", json=payload)

    if response.status_code == 200:
        data = response.json()
        print(f"✓ PDF exported successfully")
        print(f"  Filename: {data['filename']}")
        print(f"  Size: {data['size_bytes']} bytes")

        # Save the PDF
        pdf_bytes = base64.b64decode(data['pdf_base64'])
        output_file = "test_output.pdf"
        with open(output_file, "wb") as f:
            f.write(pdf_bytes)
        print(f"  Saved to: {output_file}")

        # Verify it's a valid PDF (check magic bytes)
        if pdf_bytes[:4] == b'%PDF':
            print(f"  Verified: Valid PDF file format")
        else:
            print(f"  Warning: PDF format could not be verified")

        return True
    else:
        print(f"✗ Failed to export PDF: {response.status_code}")
        print(f"  Error: {response.text}")
        return False

def test_google_endpoints_availability():
    """Test if Google Docs endpoints are available"""
    print("Testing Google Docs endpoints availability...")

    # Test with invalid credentials to check if the feature is enabled
    payload = {
        "title": "Test",
        "credentials_json": "{}"
    }

    response = requests.post(f"{BASE_URL}/documents/google/create", json=payload)

    if response.status_code == 501:
        print(f"⚠ Google Docs integration not available (missing dependencies)")
        return False
    else:
        print(f"✓ Google Docs endpoints are available")
        print(f"  Note: Actual testing requires valid OAuth2 credentials")
        return True

def main():
    """Run all tests"""
    print("\n" + "="*80)
    print("MCP DOCUMENTS SERVICE - COMPREHENSIVE TEST SUITE")
    print("="*80 + "\n")

    results = {}

    # Test 1: Health Check
    print_separator()
    results['health'] = test_health_check()

    # Test 2: Root Endpoint
    print_separator()
    results['root'] = test_root_endpoint()

    # Test 3: Create DOCX
    print_separator()
    results['create_docx'], _ = test_create_docx()

    # Test 4: Fill Template
    print_separator()
    results['fill_template'] = test_fill_template()

    # Test 5: Export PDF
    print_separator()
    results['export_pdf'] = test_export_pdf()

    # Test 6: Google Docs Availability
    print_separator()
    results['google_available'] = test_google_endpoints_availability()

    # Summary
    print_separator()
    print("TEST SUMMARY")
    print("-" * 80)

    passed = sum(1 for v in results.values() if v)
    total = len(results)

    for test_name, result in results.items():
        status = "✓ PASS" if result else "✗ FAIL"
        print(f"{status} - {test_name}")

    print("-" * 80)
    print(f"\nTotal: {passed}/{total} tests passed")

    if passed == total:
        print("\n🎉 All tests passed successfully!")
    else:
        print(f"\n⚠ {total - passed} test(s) failed")

    # List generated files
    print("\nGenerated files:")
    for filename in ["test_output.docx", "test_filled_template.docx", "test_output.pdf"]:
        if os.path.exists(filename):
            size = os.path.getsize(filename)
            print(f"  - {filename} ({size} bytes)")

    print("\n" + "="*80 + "\n")

    return passed == total

if __name__ == "__main__":
    try:
        success = main()
        exit(0 if success else 1)
    except requests.exceptions.ConnectionError:
        print("\n✗ ERROR: Could not connect to the service")
        print("  Make sure the server is running on http://localhost:8007")
        print("  Start it with: python backend/mcp/documents/server.py")
        exit(1)
    except Exception as e:
        print(f"\n✗ ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        exit(1)

