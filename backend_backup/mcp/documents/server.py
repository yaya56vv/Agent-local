"""
MCP Documents Service
Handles Word documents (DOCX), PDF export, and Google Docs integration.
"""

from fastapi import FastAPI, HTTPException
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field
from typing import Optional, Dict, Any, List
import base64
import io
import logging
from datetime import datetime

# Word documents
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# PDF export
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER, TA_RIGHT

# Google Docs (optional, requires authentication)
try:
    from google.oauth2.credentials import Credentials
    from googleapiclient.discovery import build
    from googleapiclient.errors import HttpError
    GOOGLE_AVAILABLE = True
except ImportError:
    GOOGLE_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize FastAPI
app = FastAPI(
    title="MCP Documents Service",
    description="Document processing service for Word, PDF, and Google Docs",
    version="1.0.0"
)


# ========== Pydantic Models ==========

class CreateDocxRequest(BaseModel):
    title: str = Field(..., description="Document title")
    content: List[Dict[str, Any]] = Field(..., description="List of paragraphs with formatting")
    # content format: [{"text": "...", "style": "Normal|Heading1|Heading2", "alignment": "left|center|right", "bold": bool, "italic": bool}]

class FillTemplateRequest(BaseModel):
    template_base64: str = Field(..., description="Base64 encoded template DOCX file")
    replacements: Dict[str, str] = Field(..., description="Key-value pairs for replacement ({{key}} -> value)")

class ExportPdfRequest(BaseModel):
    title: str = Field(..., description="Document title")
    content: List[Dict[str, Any]] = Field(..., description="List of paragraphs with formatting")
    page_size: str = Field(default="letter", description="Page size: letter or A4")

class GoogleCreateRequest(BaseModel):
    title: str = Field(..., description="Document title")
    credentials_json: str = Field(..., description="Google OAuth2 credentials as JSON string")

class GoogleUpdateRequest(BaseModel):
    document_id: str = Field(..., description="Google Docs document ID")
    content: str = Field(..., description="Content to set/replace")
    credentials_json: str = Field(..., description="Google OAuth2 credentials as JSON string")

class GoogleAppendRequest(BaseModel):
    document_id: str = Field(..., description="Google Docs document ID")
    content: str = Field(..., description="Content to append")
    credentials_json: str = Field(..., description="Google OAuth2 credentials as JSON string")


# ========== Helper Functions ==========

def docx_to_base64(doc: Document) -> str:
    """Convert a Document object to base64 string."""
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return base64.b64encode(buffer.read()).decode('utf-8')

def base64_to_docx(base64_string: str) -> Document:
    """Convert base64 string to Document object."""
    docx_bytes = base64.b64decode(base64_string)
    buffer = io.BytesIO(docx_bytes)
    return Document(buffer)

def pdf_to_base64(pdf_buffer: io.BytesIO) -> str:
    """Convert PDF buffer to base64 string."""
    pdf_buffer.seek(0)
    return base64.b64encode(pdf_buffer.read()).decode('utf-8')


# ========== Word Document Endpoints ==========

@app.post("/documents/create_docx")
async def create_docx(request: CreateDocxRequest):
    """
    Create a new Word document from scratch.

    Returns the document as base64 encoded string.
    """
    try:
        doc = Document()

        # Add title
        title_paragraph = doc.add_heading(request.title, level=0)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add content paragraphs
        for item in request.content:
            text = item.get("text", "")
            style = item.get("style", "Normal")
            alignment = item.get("alignment", "left")
            bold = item.get("bold", False)
            italic = item.get("italic", False)
            font_size = item.get("font_size", None)

            # Handle heading styles
            if style.startswith("Heading"):
                level = int(style.replace("Heading", "")) if style != "Heading" else 1
                paragraph = doc.add_heading(text, level=level)
            else:
                paragraph = doc.add_paragraph(text)

            # Apply alignment
            alignment_map = {
                "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            }
            paragraph.alignment = alignment_map.get(alignment, WD_PARAGRAPH_ALIGNMENT.LEFT)

            # Apply formatting to runs
            if style == "Normal" or not style.startswith("Heading"):
                for run in paragraph.runs:
                    run.bold = bold
                    run.italic = italic
                    if font_size:
                        run.font.size = Pt(font_size)

        # Convert to base64
        doc_base64 = docx_to_base64(doc)

        logger.info(f"Created DOCX document: {request.title}")
        return JSONResponse(content={
            "success": True,
            "document_base64": doc_base64,
            "filename": f"{request.title}.docx",
            "size_bytes": len(base64.b64decode(doc_base64))
        })

    except Exception as e:
        logger.error(f"Error creating DOCX: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to create document: {str(e)}")


@app.post("/documents/fill_docx_template")
async def fill_docx_template(request: FillTemplateRequest):
    """
    Fill a Word template with replacements.

    Replaces all occurrences of {{key}} with corresponding values.
    """
    try:
        # Load template from base64
        doc = base64_to_docx(request.template_base64)

        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in request.replacements.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    # Replace in inline text
                    inline = paragraph.runs
                    for run in inline:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in request.replacements.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in paragraph.text:
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, value)

        # Convert to base64
        doc_base64 = docx_to_base64(doc)

        logger.info(f"Filled DOCX template with {len(request.replacements)} replacements")
        return JSONResponse(content={
            "success": True,
            "document_base64": doc_base64,
            "filename": "filled_template.docx",
            "replacements_count": len(request.replacements)
        })

    except Exception as e:
        logger.error(f"Error filling template: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fill template: {str(e)}")


@app.post("/documents/export_pdf")
async def export_pdf(request: ExportPdfRequest):
    """
    Export content to PDF format.

    Returns the PDF as base64 encoded string.
    """
    try:
        buffer = io.BytesIO()

        # Set page size
        page_size = A4 if request.page_size.lower() == "a4" else letter
        doc = SimpleDocTemplate(buffer, pagesize=page_size,
                               rightMargin=72, leftMargin=72,
                               topMargin=72, bottomMargin=18)

        # Container for the 'Flowable' objects
        elements = []

        # Define styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='Justify', alignment=TA_JUSTIFY))

        # Add title
        title_style = styles['Title']
        elements.append(Paragraph(request.title, title_style))
        elements.append(Spacer(1, 12))

        # Add content
        for item in request.content:
            text = item.get("text", "")
            style_name = item.get("style", "Normal")
            alignment = item.get("alignment", "left")

            # Map alignment
            alignment_map = {
                "left": TA_LEFT,
                "center": TA_CENTER,
                "right": TA_RIGHT,
                "justify": TA_JUSTIFY
            }

            # Get base style
            if style_name.startswith("Heading"):
                level = style_name.replace("Heading", "")
                if level == "1":
                    style = styles['Heading1']
                elif level == "2":
                    style = styles['Heading2']
                else:
                    style = styles['Heading3']
            else:
                style = styles['Normal']

            # Create custom style with alignment
            custom_style = ParagraphStyle(
                name=f'Custom_{alignment}',
                parent=style,
                alignment=alignment_map.get(alignment, TA_LEFT)
            )

            elements.append(Paragraph(text, custom_style))
            elements.append(Spacer(1, 12))

        # Build PDF
        doc.build(elements)

        # Convert to base64
        pdf_base64 = pdf_to_base64(buffer)

        logger.info(f"Created PDF document: {request.title}")
        return JSONResponse(content={
            "success": True,
            "pdf_base64": pdf_base64,
            "filename": f"{request.title}.pdf",
            "size_bytes": len(base64.b64decode(pdf_base64))
        })

    except Exception as e:
        logger.error(f"Error exporting PDF: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to export PDF: {str(e)}")


# ========== Google Docs Endpoints ==========

@app.post("/documents/google/create")
async def google_create(request: GoogleCreateRequest):
    """
    Create a new Google Docs document.

    Requires valid Google OAuth2 credentials.
    """
    if not GOOGLE_AVAILABLE:
        raise HTTPException(
            status_code=501,
            detail="Google Docs integration not available. Install google-api-python-client."
        )

    try:
        import json

        # Parse credentials
        creds_dict = json.loads(request.credentials_json)
        creds = Credentials.from_authorized_user_info(creds_dict)

        # Build the service
        service = build('docs', 'v1', credentials=creds)

        # Create document
        doc_body = {'title': request.title}
        doc = service.documents().create(body=doc_body).execute()

        document_id = doc.get('documentId')

        logger.info(f"Created Google Doc: {request.title} (ID: {document_id})")
        return JSONResponse(content={
            "success": True,
            "document_id": document_id,
            "title": request.title,
            "url": f"https://docs.google.com/document/d/{document_id}/edit"
        })

    except HttpError as e:
        logger.error(f"Google API error: {str(e)}")
        raise HTTPException(status_code=502, detail=f"Google API error: {str(e)}")
    except Exception as e:
        logger.error(f"Error creating Google Doc: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to create Google Doc: {str(e)}")


@app.post("/documents/google/update")
async def google_update(request: GoogleUpdateRequest):
    """
    Update/replace content in a Google Docs document.

    Replaces all existing content with new content.
    """
    if not GOOGLE_AVAILABLE:
        raise HTTPException(
            status_code=501,
            detail="Google Docs integration not available. Install google-api-python-client."
        )

    try:
        import json

        # Parse credentials
        creds_dict = json.loads(request.credentials_json)
        creds = Credentials.from_authorized_user_info(creds_dict)

        # Build the service
        service = build('docs', 'v1', credentials=creds)

        # Get document to find the end index
        doc = service.documents().get(documentId=request.document_id).execute()
        end_index = doc.get('body').get('content')[-1].get('endIndex') - 1

        # Delete all content and insert new
        requests = [
            {
                'deleteContentRange': {
                    'range': {
                        'startIndex': 1,
                        'endIndex': end_index
                    }
                }
            },
            {
                'insertText': {
                    'location': {
                        'index': 1
                    },
                    'text': request.content
                }
            }
        ]

        result = service.documents().batchUpdate(
            documentId=request.document_id,
            body={'requests': requests}
        ).execute()

        logger.info(f"Updated Google Doc: {request.document_id}")
        return JSONResponse(content={
            "success": True,
            "document_id": request.document_id,
            "url": f"https://docs.google.com/document/d/{request.document_id}/edit"
        })

    except HttpError as e:
        logger.error(f"Google API error: {str(e)}")
        raise HTTPException(status_code=502, detail=f"Google API error: {str(e)}")
    except Exception as e:
        logger.error(f"Error updating Google Doc: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to update Google Doc: {str(e)}")


@app.post("/documents/google/append")
async def google_append(request: GoogleAppendRequest):
    """
    Append content to a Google Docs document.

    Adds content at the end of the document.
    """
    if not GOOGLE_AVAILABLE:
        raise HTTPException(
            status_code=501,
            detail="Google Docs integration not available. Install google-api-python-client."
        )

    try:
        import json

        # Parse credentials
        creds_dict = json.loads(request.credentials_json)
        creds = Credentials.from_authorized_user_info(creds_dict)

        # Build the service
        service = build('docs', 'v1', credentials=creds)

        # Get document to find the end index
        doc = service.documents().get(documentId=request.document_id).execute()
        end_index = doc.get('body').get('content')[-1].get('endIndex') - 1

        # Insert at the end
        requests = [
            {
                'insertText': {
                    'location': {
                        'index': end_index
                    },
                    'text': '\n' + request.content
                }
            }
        ]

        result = service.documents().batchUpdate(
            documentId=request.document_id,
            body={'requests': requests}
        ).execute()

        logger.info(f"Appended to Google Doc: {request.document_id}")
        return JSONResponse(content={
            "success": True,
            "document_id": request.document_id,
            "url": f"https://docs.google.com/document/d/{request.document_id}/edit"
        })

    except HttpError as e:
        logger.error(f"Google API error: {str(e)}")
        raise HTTPException(status_code=502, detail=f"Google API error: {str(e)}")
    except Exception as e:
        logger.error(f"Error appending to Google Doc: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to append to Google Doc: {str(e)}")


# ========== Health Check ==========

@app.get("/health")
async def health_check():
    """Health check endpoint."""
    return {
        "status": "healthy",
        "service": "mcp-documents",
        "timestamp": datetime.now().isoformat(),
        "google_docs_available": GOOGLE_AVAILABLE
    }


@app.get("/")
async def root():
    """Root endpoint with service information."""
    return {
        "service": "MCP Documents Service",
        "version": "1.0.0",
        "endpoints": {
            "word": [
                "POST /documents/create_docx",
                "POST /documents/fill_docx_template"
            ],
            "pdf": [
                "POST /documents/export_pdf"
            ],
            "google": [
                "POST /documents/google/create",
                "POST /documents/google/update",
                "POST /documents/google/append"
            ]
        },
        "google_docs_available": GOOGLE_AVAILABLE
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8007)

